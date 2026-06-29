import re
from io import BytesIO
from pathlib import Path

from docx import Document, section


TEMPLATE_PATH = Path(__file__).resolve().parents[2] / "Sample.docx"


def _validate_content(name: str, value: str) -> str:
    if not isinstance(value, str):
        raise TypeError(f"{name} must be a string.")
    return value.strip()


def _extract_section(
    text: str,
    heading: str,
    following_headings: tuple[str, ...],
) -> str:
    if not following_headings:
        match = re.search(
            rf"(?ims)^\s*{re.escape(heading)}\s*:?\s*$\s*(.*)\Z",
            text,
        )
        return match.group(1).strip() if match else ""

    end_pattern = "|".join(re.escape(item) for item in following_headings)
    match = re.search(
        rf"(?ims)^\s*{re.escape(heading)}\s*:?\s*$\s*(.*?)"
        rf"(?=^\s*(?:{end_pattern})\s*:?\s*$|\Z)",
        text,
    )
    return match.group(1).strip() if match else ""

def _parse_why_section(section):

    section = section.strip()

    if not section:
        return "", ""

    lines = [x.strip() for x in section.splitlines() if x.strip()]

    question = ""
    answer = ""

    for i, line in enumerate(lines):

        if line.lower().startswith("answer:"):

            answer = line[7:].strip()

            if i + 1 < len(lines):
                answer += "\n" + "\n".join(lines[i + 1:])

            break

        elif not question:

            question = line

    return question, answer
def _parse_incident(incident_text: str) -> dict[str, str]:
    headings = ("Problem Description", "Business Impact", "Investigation Timeline")
    return {
        heading: _extract_section(incident_text, heading, headings[index + 1 :])
        for index, heading in enumerate(headings)
    }


def _replace_paragraph_text(paragraph, value: str) -> None:
    for run in paragraph.runs:
        run.text = ""

    if not paragraph.runs:
        paragraph.add_run()

    lines = value.splitlines()
    if not lines:
        return

    paragraph.runs[0].text = lines[0]
    for line in lines[1:]:
        paragraph.runs[0].add_break()
        paragraph.runs[0].add_text(line)


def _replace_cell_text(cell, value: str) -> None:
    first_paragraph = cell.paragraphs[0]
    for paragraph in cell.paragraphs[1:]:
        cell._tc.remove(paragraph._p)
    _replace_paragraph_text(first_paragraph, value)


def _clear_table_body(table, start_row: int = 1) -> None:
    for row in table.rows[start_row:]:
        for cell in row.cells:
            _replace_cell_text(cell, "")


def _timeline_entries(timeline: str, capacity: int) -> list[tuple[str, str, str]]:
    lines = [line.strip() for line in timeline.splitlines() if line.strip()]
    entries = []

    for line in lines:
        date_value = ""
        time_value = ""
        activity = line

        date_match = re.match(
            r"^(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})\s+(.+)$",
            activity,
        )
        if date_match:
            date_value = date_match.group(1)
            activity = date_match.group(2).strip()

        time_match = re.match(
            r"^(\d{1,2}:\d{2}(?:\s*[AP]M)?)\s*[-–:]?\s*(.*)$",
            activity,
            flags=re.IGNORECASE,
        )
        if time_match:
            time_value = time_match.group(1)
            activity = time_match.group(2).strip()

        entries.append((date_value, time_value, activity))

    if len(entries) > capacity:
        overflow = "\n".join(item[2] for item in entries[capacity - 1 :])
        entries = entries[: capacity - 1] + [("", "", overflow)]

    return entries


def _populate_template(
    document: Document,
    problem_description: str,
    business_impact: str,
    timeline: str,
    rca_text: str,
    capa_text: str,
) -> None:
    if len(document.tables) < 8:
        raise ValueError("Sample.docx does not contain the expected RCA-CAPA tables.")

    prepared_by_table = document.tables[0]
    for row in prepared_by_table.rows[1:]:
        _replace_cell_text(row.cells[1], "")

    problem_table = document.tables[1]
    for row in problem_table.rows:
        for cell in row.cells[1:]:
            _replace_cell_text(cell, "")
    _replace_cell_text(problem_table.cell(1, 1), problem_description)
    _replace_cell_text(problem_table.cell(2, 1), business_impact)

    related_tickets_table = document.tables[2]
    _clear_table_body(related_tickets_table)

    timeline_table = document.tables[3]
    _clear_table_body(timeline_table, start_row=2)
    timeline_rows = timeline_table.rows[2:]
    for row, entry in zip(
        timeline_rows,
        _timeline_entries(timeline, len(timeline_rows)),
    ):
        date_value, time_value, activity = entry
        _replace_cell_text(row.cells[0], date_value)
        _replace_cell_text(row.cells[1], time_value)
        _replace_cell_text(row.cells[2], activity)

    why_table = document.tables[4]
    _clear_table_body(why_table)
    why_headings = tuple(f"Why {index}" for index in range(1, 6))
    for index, row in enumerate(why_table.rows[1:], start=1):
        following = tuple(f"Why {item}" for item in range(index + 1, 6))
        section = _extract_section(
            rca_text,
            f"Why {index}",
            following,
        )

        question, answer = _parse_why_section(section)

        _replace_cell_text(
            row.cells[0],
            question,
        )

        _replace_cell_text(
            row.cells[1],
            answer,
        )

    root_cause_table = document.tables[5]
    _clear_table_body(root_cause_table)
    root_cause = _extract_section(rca_text, "Root Cause", why_headings)
    _replace_cell_text(root_cause_table.cell(1, 0), root_cause)

    actions_table = document.tables[6]
    _clear_table_body(actions_table)
    capa_headings = (
        "Corrective Action",
        "Preventive Action",
        "Ownership",
        "Monitoring",
        "Lessons Learned",
    )
    corrective_action = _extract_section(
        capa_text,
        "Corrective Action",
        capa_headings[1:],
    )
    preventive_action = _extract_section(
        capa_text,
        "Preventive Action",
        capa_headings[2:],
    )
    _replace_cell_text(actions_table.cell(1, 0), "CA")
    _replace_cell_text(actions_table.cell(1, 1), corrective_action)
    _replace_cell_text(actions_table.cell(2, 0), "PA")
    _replace_cell_text(actions_table.cell(2, 1), preventive_action)

    lessons_learned = _extract_section(
        capa_text,
        "Lessons Learned",
        (),
    )
    lesson_heading_index = next(
        (
            index
            for index, paragraph in enumerate(document.paragraphs)
            if paragraph.text.strip() == "Lessons Learned"
        ),
        None,
    )
    if lesson_heading_index is not None:
        for paragraph in document.paragraphs[lesson_heading_index + 1 :]:
            if paragraph.style.name.startswith("Heading"):
                break
            if paragraph.text.strip():
                _replace_paragraph_text(paragraph, lessons_learned)
                lessons_learned = ""

    references_heading_index = next(
        (
            index
            for index, paragraph in enumerate(document.paragraphs)
            if paragraph.text.strip() == "References"
        ),
        None,
    )
    if references_heading_index is not None:
        for paragraph in document.paragraphs[references_heading_index + 1 :]:
            if paragraph.style.name.startswith("Heading"):
                break
            _replace_paragraph_text(paragraph, "")

    version_history_table = document.tables[7]
    _clear_table_body(version_history_table)


def generate_docx(
    incident_text: str,
    rca_text: str,
    capa_text: str,
    problem_description: str = "",
    business_impact: str = "",
    investigation_timeline: str = "",
) -> BytesIO:
    incident = _validate_content("incident_text", incident_text)
    rca = _validate_content("rca_text", rca_text)
    capa = _validate_content("capa_text", capa_text)

    if not TEMPLATE_PATH.is_file():
        raise FileNotFoundError(f"DOCX template not found: {TEMPLATE_PATH}")

    parsed_incident = _parse_incident(incident)
    problem = _validate_content("problem_description", problem_description)
    impact = _validate_content("business_impact", business_impact)
    timeline = _validate_content("investigation_timeline", investigation_timeline)

    document = Document(str(TEMPLATE_PATH))
    _populate_template(
        document=document,
        problem_description=problem or parsed_incident["Problem Description"],
        business_impact=impact or parsed_incident["Business Impact"],
        timeline=timeline or parsed_incident["Investigation Timeline"],
        rca_text=rca,
        capa_text=capa,
    )

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output
